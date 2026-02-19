# event_resume_parser.py
# -*- coding: utf-8 -*-
"""
Event Resume Parser
-------------------
Given a file (or folder) of event resumes, extract:
  • Main arrival date
  • Main departure date
  • Fly-in percentage
  • Drive-in percentage
  • Valet prices (if any)

Supported inputs: PDF (.pdf), Word (.docx), Text (.txt), Excel/CSV (.xlsx, .xls, .csv)

Usage examples:
  python event_resume_parser.py --input "/path/to/resumes" --output summary.csv
  python event_resume_parser.py --input "/path/to/file.pdf" --print

Notes & assumptions:
  • No OCR: scanned images inside PDFs are not read.
  • If multiple dates are mentioned, the script will select the earliest as main arrival and the latest as main departure.
  • Percentages are not inferred; only what’s explicitly stated near relevant keywords is returned.
  • For ambiguous formats, improve results by adding more keywords to the lists below.
"""

import argparse
import os
import re
import sys
from datetime import datetime
from typing import Dict, List, Optional, Tuple

# Optional imports gated with try/except so the script still runs without some libs
try:
    from PyPDF2 import PdfReader
except Exception:  # pragma: no cover
    PdfReader = None  # type: ignore

try:
    import docx
except Exception:  # pragma: no cover
    docx = None  # type: ignore

try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None  # type: ignore

# -----------------------------
# Text extraction per file type
# -----------------------------

def extract_text_from_pdf(path: str) -> str:
    if PdfReader is None:
        return ""
    try:
        reader = PdfReader(path)
        texts = []
        for page in reader.pages:
            try:
                texts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(texts)
    except Exception:
        return ""


def extract_text_from_docx(path: str) -> str:
    if docx is None:
        return ""
    try:
        d = docx.Document(path)
        parts: List[str] = []
        # Paragraphs
        for p in d.paragraphs:
            parts.append(p.text)
        # Tables
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_txt(path: str) -> str:
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def extract_text_from_spreadsheet(path: str) -> str:
    if pd is None:
        return ""
    try:
        text_chunks: List[str] = []
        _, ext = os.path.splitext(path.lower())
        if ext == ".csv":
            try:
                df = pd.read_csv(path)
                text_chunks.append(df.astype(str).to_string(index=False))
            except Exception:
                pass
        else:  # .xlsx or .xls
            try:
                engine = "openpyxl" if ext == ".xlsx" else "xlrd"
                xl = pd.ExcelFile(path, engine=engine)
                for sheet in xl.sheet_names:
                    df = xl.parse(sheet)
                    text_chunks.append(df.astype(str).to_string(index=False))
            except Exception:
                pass
        return "\n".join(text_chunks)
    except Exception:
        return ""


EXT_READERS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_txt,
    ".xlsx": extract_text_from_spreadsheet,
    ".xls": extract_text_from_spreadsheet,
    ".csv": extract_text_from_spreadsheet,
}


# -----------------------------
# Heuristics & regexes
# -----------------------------

MONTHS = (
    "january", "february", "march", "april", "may", "june",
    "july", "august", "september", "october", "november", "december",
)
MONTH_ABBR = (
    "jan", "feb", "mar", "apr", "may", "jun",
    "jul", "aug", "sep", "sept", "oct", "nov", "dec",
)

# Regex for textual month dates, e.g., "January 3, 2026" or "Jan 3-5, 2026"
TEXTUAL_DATE_RE = re.compile(
    r"\b(?:(?:mon|tue|wed|thu|thur|thurs|fri|sat|sun)\.?\s*)?"  # optional weekday
    r"((?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|"
    r"sep(?:t|tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?))\s+"
    r"(\d{1,2})(?:st|nd|rd|th)?"                                  # day
    r"(?:\s*[-–—]\s*(\d{1,2})(?:st|nd|rd|th)?)?"                 # optional end day (range)
    r"(?:,?\s*(\d{4}))?\b",                                      # optional year
    re.IGNORECASE,
)

# Regex for numeric dates like 1/3/2026, 01-03-26
NUMERIC_DATE_RE = re.compile(r"\b(\d{1,2})[/-](\d{1,2})(?:[/-](\d{2,4}))\b")

# Keywords to anchor arrival/departure cues
ARRIVAL_KEYS = [
    "arrival", "arrive", "check-in", "check in", "inbound", "main arrival",
]
DEPARTURE_KEYS = [
    "departure", "depart", "check-out", "check out", "outbound", "main departure",
]

# Keywords for transport modes
FLY_KEYS = ["fly", "fly-in", "fly in", "air", "airlift", "flight", "plane"]
DRIVE_KEYS = ["drive", "drive-in", "drive in", "ground", "car", "local", "bus", "motorcoach"]

PERCENT_RE = re.compile(r"(\d{1,3})\s*%")

# Valet price patterns
CURRENCY_RE = re.compile(r"\$\s*(\d+(?:\.\d{1,2})?)\s*(\+\+)?")
VALET_KEYS = ["valet", "valet parking", "overnight valet", "day valet"]


def normalize_whitespace(text: str) -> str:
    return re.sub(r"[\u00A0\s]+", " ", text).strip()


def split_sentences(text: str) -> List[str]:
    # Simple sentence splitter to keep nearby context
    parts = re.split(r"(?<=[\.!?])\s+|\n+", text)
    return [p.strip() for p in parts if p.strip()]


# -----------------------------
# Date parsing helpers
# -----------------------------

MONTH_NAME_TO_NUM: Dict[str, int] = {
    "jan": 1, "january": 1,
    "feb": 2, "february": 2,
    "mar": 3, "march": 3,
    "apr": 4, "april": 4,
    "may": 5,
    "jun": 6, "june": 6,
    "jul": 7, "july": 7,
    "aug": 8, "august": 8,
    "sep": 9, "sept": 9, "september": 9,
    "oct": 10, "october": 10,
    "nov": 11, "november": 11,
    "dec": 12, "december": 12,
}


def try_parse_year_from_text(text: str) -> Optional[int]:
    years = re.findall(r"\b(20\d{2}|19\d{2})\b", text)
    if years:
        try:
            return int(years[0])
        except Exception:
            return None
    return None


def make_date(year: int, month: int, day: int) -> Optional[datetime]:
    try:
        return datetime(year, month, day)
    except Exception:
        return None


def parse_textual_dates(text: str, fallback_year: Optional[int]) -> List[Tuple[Optional[datetime], Optional[datetime]]]:
    """Return list of (start_date, end_date) where each element may be None if unknown."""
    results: List[Tuple[Optional[datetime], Optional[datetime]]] = []
    for m in TEXTUAL_DATE_RE.finditer(text):
        mon_str = m.group(1)
        day1 = m.group(2)
        day2 = m.group(3)
        year = m.group(4)
        month_num = MONTH_NAME_TO_NUM.get(mon_str.lower(), None) if mon_str else None
        year_num: Optional[int] = int(year) if year else (fallback_year if fallback_year else None)
        if month_num is None:
            continue
        try:
            d1 = int(day1)
        except Exception:
            d1 = None  # type: ignore
        d2 = int(day2) if day2 else None
        if year_num is None or d1 is None:
            results.append((None, None))
            continue
        start = make_date(year_num, month_num, d1)
        end = make_date(year_num, month_num, d2) if d2 else None
        results.append((start, end))
    return results


def parse_numeric_dates(text: str) -> List[datetime]:
    dates: List[datetime] = []
    for m in NUMERIC_DATE_RE.finditer(text):
        mth, day, yr = m.group(1), m.group(2), m.group(3)
        try:
            mm = int(mth)
            dd = int(day)
            yy = int(yr)
            if yy < 100:
                # pivot: 00-69 -> 2000-2069, 70-99 -> 1970-1999
                yy = 2000 + yy if yy <= 69 else 1900 + yy
            dt = make_date(yy, mm, dd)
            if dt:
                dates.append(dt)
        except Exception:
            continue
    return dates


def choose_main_arrival_departure(text: str) -> Tuple[Optional[datetime], Optional[datetime]]:
    """Heuristic: use earliest arrival-like date as arrival; latest departure-like date as departure.
    If labeled ranges are found (e.g., Jan 3–5), use their bounds as candidates.
    """
    fallback_year = try_parse_year_from_text(text)
    sentences = split_sentences(text)

    arrival_candidates: List[datetime] = []
    departure_candidates: List[datetime] = []

    # Labeled by keywords
    for sent in sentences:
        sent_low = sent.lower()
        labeled_arrival = any(k in sent_low for k in ARRIVAL_KEYS)
        labeled_depart = any(k in sent_low for k in DEPARTURE_KEYS)

        # Textual ranges within sentence
        for start, end in parse_textual_dates(sent, fallback_year):
            if start and labeled_arrival:
                arrival_candidates.append(start)
            if end and labeled_depart:
                departure_candidates.append(end)
            # If sentence has both labels (or none) still consider starts/ends as generic candidates
            if start and not labeled_arrival and not labeled_depart:
                arrival_candidates.append(start)
            if end and not labeled_depart and not labeled_arrival:
                departure_candidates.append(end)

        # Numeric dates within sentence
        for dt in parse_numeric_dates(sent):
            if labeled_arrival:
                arrival_candidates.append(dt)
            if labeled_depart:
                departure_candidates.append(dt)

    # Global scan if nothing found via labeled sentences
    if not arrival_candidates or not departure_candidates:
        global_ranges = parse_textual_dates(text, fallback_year)
        for start, end in global_ranges:
            if start:
                arrival_candidates.append(start)
            if end:
                departure_candidates.append(end)
        for dt in parse_numeric_dates(text):
            arrival_candidates.append(dt)
            departure_candidates.append(dt)

    arrival = min(arrival_candidates) if arrival_candidates else None
    departure = max(departure_candidates) if departure_candidates else None

    # Sanity: if arrival after departure, swap or null out
    if arrival and departure and arrival > departure:
        # As last resort, null departure
        departure = None
    return arrival, departure


# -----------------------------
# Percentage & valet extraction
# -----------------------------

def window_contains_any(words: List[str], keys: List[str]) -> bool:
    text = " ".join(words).lower()
    return any(k in text for k in keys)


def extract_transport_percentages(text: str) -> Tuple[Optional[int], Optional[int]]:
    """Return (fly_in_pct, drive_in_pct). Only values explicitly tied to relevant keywords are returned."""
    fly_pct: Optional[int] = None
    drive_pct: Optional[int] = None
    tokens = re.split(r"\s+", text)

    # Map indices of % to nearest mode
    for i, tok in enumerate(tokens):
        m = PERCENT_RE.match(tok)
        if m:
            try:
                val = int(m.group(1))
            except Exception:
                continue
            # Look around a window
            left = max(0, i - 6)
            right = min(len(tokens), i + 7)
            window = tokens[left:right]
            if window_contains_any(window, FLY_KEYS):
                fly_pct = val if fly_pct is None else fly_pct
            if window_contains_any(window, DRIVE_KEYS):
                drive_pct = val if drive_pct is None else drive_pct

    # Fallback: sentence-based proximity
    if fly_pct is None or drive_pct is None:
        for sent in split_sentences(text):
            if (fly_pct is None) and any(k in sent.lower() for k in FLY_KEYS):
                m = PERCENT_RE.search(sent)
                if m:
                    try:
                        fly_pct = int(m.group(1))
                    except Exception:
                        pass
            if (drive_pct is None) and any(k in sent.lower() for k in DRIVE_KEYS):
                m = PERCENT_RE.search(sent)
                if m:
                    try:
                        drive_pct = int(m.group(1))
                    except Exception:
                        pass
    return fly_pct, drive_pct


def extract_valet_prices(text: str) -> Tuple[Optional[float], Optional[str]]:
    """Return the first detected valet price and its source sentence as notes."""
    for sent in split_sentences(text):
        low = sent.lower()
        if any(k in low for k in VALET_KEYS):
            for m in CURRENCY_RE.finditer(sent):
                try:
                    amt = float(m.group(1))
                except Exception:
                    continue
                notes = sent.strip()
                return amt, notes
    # Broader search: any sentence with 'parking' and a price mentioning valet in nearby words
    for sent in split_sentences(text):
        low = sent.lower()
        if "parking" in low:
            for m in CURRENCY_RE.finditer(sent):
                try:
                    amt = float(m.group(1))
                except Exception:
                    continue
                notes = sent.strip()
                return amt, notes
    return None, None


# -----------------------------
# Core processing
# -----------------------------

def extract_from_text(raw_text: str) -> Dict[str, Optional[str]]:
    text = normalize_whitespace(raw_text)
    arr, dep = choose_main_arrival_departure(text)
    fly, drive = extract_transport_percentages(text)
    valet_amt, valet_notes = extract_valet_prices(text)

    def fmt(dt: Optional[datetime]) -> Optional[str]:
        return dt.strftime("%Y-%m-%d") if dt else None

    result: Dict[str, Optional[str]] = {
        "arrival_date": fmt(arr),
        "departure_date": fmt(dep),
        "fly_in_pct": str(fly) if fly is not None else None,
        "drive_in_pct": str(drive) if drive is not None else None,
        "valet_price": (f"{valet_amt:.2f}" if valet_amt is not None else None),
        "valet_notes": valet_notes,
    }
    return result


def extract_from_file(path: str) -> Dict[str, Optional[str]]:
    ext = os.path.splitext(path.lower())[1]
    reader = EXT_READERS.get(ext)
    text = ""
    if reader:
        text = reader(path)
    # As a fallback, try reading as text
    if not text:
        text = extract_text_from_txt(path)
    data = extract_from_text(text)
    data["source_file"] = os.path.basename(path)
    return data


def walk_files(input_path: str) -> List[str]:
    if os.path.isdir(input_path):
        files: List[str] = []
        for root, _, names in os.walk(input_path):
            for n in names:
                ext = os.path.splitext(n.lower())[1]
                if ext in EXT_READERS:
                    files.append(os.path.join(root, n))
        return files
    else:
        return [input_path]


def to_csv(rows: List[Dict[str, Optional[str]]], out_path: str) -> None:
    headers = [
        "source_file",
        "arrival_date",
        "departure_date",
        "fly_in_pct",
        "drive_in_pct",
        "valet_price",
        "valet_notes",
    ]
    # Write CSV without pandas for fewer dependencies
    import csv
    with open(out_path, "w", encoding="utf-8", newline="") as f:
        w = csv.DictWriter(f, fieldnames=headers)
        w.writeheader()
        for r in rows:
            w.writerow({k: r.get(k) if r.get(k) is not None else "" for k in headers})


def print_table(rows: List[Dict[str, Optional[str]]]) -> None:
    # Minimal pretty table
    headers = [
        "source_file",
        "arrival_date",
        "departure_date",
        "fly_in_pct",
        "drive_in_pct",
        "valet_price",
    ]
    col_widths = {h: max(len(h), *(len(str(r.get(h, "") or "")) for r in rows)) for h in headers}
    sep = "+" + "+".join("-" * (col_widths[h] + 2) for h in headers) + "+"
    def row_line(vals: List[str]) -> str:
        return "| " + " | ".join(v.ljust(col_widths[h]) for v, h in zip(vals, headers)) + " |"

    print(sep)
    print(row_line(headers))
    print(sep)
    for r in rows:
        vals = [
            str(r.get("source_file", "") or ""),
            str(r.get("arrival_date", "") or ""),
            str(r.get("departure_date", "") or ""),
            str(r.get("fly_in_pct", "") or ""),
            str(r.get("drive_in_pct", "") or ""),
            str(r.get("valet_price", "") or ""),
        ]
        print(row_line(vals))
    print(sep)


def main():
    parser = argparse.ArgumentParser(description="Extract key details from event resumes.")
    parser.add_argument("--input", required=True, help="Path to a file or a folder containing resumes")
    parser.add_argument("--output", help="Optional path to write a CSV summary (e.g., summary.csv)")
    parser.add_argument("--print", dest="do_print", action="store_true", help="Print a table to stdout")
    parser.add_argument("--latest-pdf", action="store_true", help="When --input is a folder, process only the most recently modified PDF in that folder")
    args = parser.parse_args()

    paths = walk_files(args.input)
    if not paths:
        print("No supported files found.")
        sys.exit(1)
        
    #If --latest-pdf and input is a directory, narrow to the most recent PDF
    if args.latest_pdf and os.path.isdir(args.input):
        pdfs = [p for p in paths if p.lower().endswith('.pdf')]    
        if not pdfs:        
            print("--latest-pdf was set but no PDF files were found in the folder.")        
            sys.exit(1)    
        latest_pdf = max(pdfs, key=lambda f: os.path.getmtime(f))    
        paths = [latest_pdf]

    rows = []
    for p in paths:
        rows.append(extract_from_file(p))

    if args.output:
        to_csv(rows, args.output)
        print(f"Wrote {len(rows)} rows to {args.output}")

    if args.do_print or not args.output:
        print_table(rows)


if __name__ == "__main__":
    main()
